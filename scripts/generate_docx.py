"""
scripts/generate_docx.py
Generates comprehensive .docx documentation for the
Offline AI Swine Health Monitoring System capstone project.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(BASE_DIR, 'docx_documentation')

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set a table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_header_row(table, headers: list, bg: str = '1F4E79'):
    """Add a bold white header row to a table."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_data_row(table, row_index: int, values: list, shade_odd: bool = True):
    """Add a data row, optionally shading odd rows."""
    row = table.rows[row_index]
    bg = 'D6E4F0' if (shade_odd and row_index % 2 == 0) else 'FFFFFF'
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        set_cell_bg(cell, bg)
        cell.paragraphs[0].runs[0].font.size = Pt(9)


def add_styled_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def add_info_box(doc, text: str, bg_hex: str = 'EBF5FB'):
    """Add a shaded paragraph box (like a note/callout)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    return p


def add_image(doc, rel_path: str, caption: str, width: float = 5.8):
    """Add an image with caption if the file exists."""
    full = os.path.join(BASE_DIR, rel_path)
    if os.path.exists(full):
        try:
            doc.add_picture(full, width=Inches(width))
            cap = doc.add_paragraph(f'Figure: {caption}')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.italic = True
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        except Exception as e:
            doc.add_paragraph(f'[Image unavailable: {rel_path} — {e}]')
    else:
        doc.add_paragraph(f'[Image file not found: {rel_path}]')


def save(doc: Document, name: str):
    path = os.path.join(OUTPUT_DIR, name)
    doc.save(path)
    print(f'  [OK]  {name}')


# ─── Document 1: General Comprehensive Overview ─────────────────────────────

def doc_general_overview():
    doc = Document()
    title = doc.add_heading('Offline AI Swine Health Monitoring System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('General Comprehensive Overview')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph('Capstone Research Project · Raspberry Pi 4B · Edge AI · Offline-First').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Executive Summary
    add_styled_heading(doc, '1. Executive Summary')
    doc.add_paragraph(
        'The Offline AI Swine Health Monitoring System is a fully autonomous, edge-computing capstone '
        'project designed to operate on a Raspberry Pi 4B without requiring cloud connectivity, internet '
        'access, or any subscription fees. It fuses computer vision, infrared thermal sensing, and ambient '
        'environmental monitoring to detect early indicators of swine illness and heat stress — delivering '
        'real-time alerts directly to the farmer via SMS.'
    )
    doc.add_paragraph(
        'Traditional swine health monitoring depends heavily on manual inspection, which is labor-intensive, '
        'inconsistent, and prone to missing early disease indicators. This system addresses that gap by '
        'continuously analyzing pig behavior and temperature around the clock, triggering automated alerts '
        'when anomalies are detected.'
    )

    # 2. Core Features
    add_styled_heading(doc, '2. Core Features & Capabilities')
    features = [
        ('Real-Time Behavior Detection',
         'A YOLOv8n model (optimized with ONNX Runtime) classifies 8 distinct pig behaviors per frame at '
         '14–30 FPS on the Raspberry Pi 4B CPU, including lying, standing, walking, sitting, feeding, '
         'drinking, social interaction, and aggression.'),
        ('Individual Pig Tracking',
         'SORT (Simple Online and Realtime Tracking) with Kalman filters assigns persistent track IDs to '
         'each detected pig, enabling per-pig behavior history and stationary duration monitoring.'),
        ('Thermal Temperature Mapping',
         'An Adafruit AMG8833 8×8 infrared thermal camera reads temperature grids at ~10 Hz. Zone '
         'centroids are matched against YOLO bounding boxes to associate body surface temperatures with '
         'individual pig IDs.'),
        ('Ambient Heat Stress Monitoring (THI)',
         'A DHT22 sensor provides real-time ambient temperature and humidity, used to compute the '
         'Temperature Humidity Index (THI). The THI dynamically adjusts alert thresholds to prevent '
         'false positives during hot weather.'),
        ('Hybrid Dual-Channel Risk Engine',
         'The system runs two parallel alert channels: Channel 1 targets individual pigs (elevated '
         'temperature + prolonged inactivity = potential fever), while Channel 2 targets herd-level '
         'lethargy (>60% of pigs stationary = potential disease outbreak).'),
        ('Offline Flask Dashboard',
         'A lightweight Flask web application serves a live dashboard accessible over the local area '
         'network (LAN) or via the Pi\'s own standalone Wi-Fi hotspot (AP mode). It displays the '
         'annotated video feed, thermal heatmap, alert history, and live statistics.'),
        ('Offline SMS Alerts',
         'A GSM900A module communicates via UART using AT commands to send SMS alerts directly to the '
         'farmer\'s mobile phone — no internet or SIM data plan required, only basic voice/SMS '
         'connectivity.'),
        ('SQLite Offline Database',
         'All detections, ambient readings, and health alerts are persisted in a local SQLite database '
         '(WAL mode) for historical analysis, audit trails, and dashboard querying.'),
    ]
    for title_text, desc in features:
        add_styled_heading(doc, f'2.{features.index((title_text, desc))+1}  {title_text}', level=2)
        doc.add_paragraph(desc)

    # 3. Tech Stack Table
    add_styled_heading(doc, '3. Technology Stack')
    doc.add_paragraph('The following table summarises every technology layer used in the system and the rationale for its selection:')

    headers = ['Layer', 'Technology', 'Rationale']
    rows = [
        ('AI Model', 'YOLOv8n (Ultralytics)', 'Smallest YOLO variant; fastest on CPU-only edge hardware'),
        ('Inference Runtime (Pi)', 'ONNX Runtime', 'Eliminates PyTorch on Pi; CPU-optimised with ORT_ENABLE_ALL'),
        ('Tracking', 'SORT — Kalman Filter + Hungarian Algorithm', 'Runs at 100+ FPS overhead; pure Python; no GPU needed'),
        ('Thermal Camera', 'Adafruit AMG8833 (8×8, I2C 0x69)', 'Offline; direct I2C GPIO; ~10 Hz refresh rate'),
        ('Ambient Sensor', 'DHT22 (GPIO4)', 'Computes THI for adaptive heat-stress thresholds'),
        ('SMS Alert', 'GSM900A (UART AT Commands)', '100% offline; no data plan; reliable hardware layer'),
        ('Database', 'SQLite 3 (WAL mode)', 'Zero-server; embedded; offline; concurrent read+write'),
        ('Dashboard', 'Flask + Jinja2 + Chart.js', 'Minimal RAM footprint; no Node.js; offline-compatible'),
        ('Networking', 'hostapd + dnsmasq (AP mode) / LAN', 'Pi creates standalone hotspot for field use'),
        ('Language', 'Python 3.11 (Bookworm)', 'Best AI/ML ecosystem; Raspberry Pi OS default'),
        ('Training Hardware', 'NVIDIA RTX 4050 + CUDA 12.x', '10–20× faster than CPU; batch=32 in 6 GB VRAM'),
    ]
    table = doc.add_table(rows=len(rows)+1, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, headers)
    for i, row in enumerate(rows):
        add_data_row(table, i+1, row)
    doc.add_paragraph()

    # 4. System Architecture
    add_styled_heading(doc, '4. System Architecture & Data Flow')
    doc.add_paragraph(
        'The system follows a modular, offline-first pipeline architecture. Each subsystem has a single '
        'responsibility and communicates with adjacent layers through well-defined interfaces.'
    )
    steps = [
        ('Frame Acquisition', 'The AsyncCamera module captures USB camera frames in a dedicated background thread, '
         'decoupling frame capture latency (0–227 ms variance) from the main inference loop.'),
        ('Thermal Acquisition', 'The ThermalReader module polls the AMG8833 sensor over I2C at ~10 Hz, retrieving '
         'the raw 8×8 float temperature grid. The grid is bilinearly scaled to 32×32 for visualization.'),
        ('Ambient Acquisition', 'The DHT22Sensor module reads temperature and humidity to compute the THI, which '
         'is logged to the database and used by the Risk Engine.'),
        ('YOLO Inference', 'The Detector module runs the ONNX YOLOv8n model at 640×640 resolution, performing '
         'preprocessing (resize + normalize), inference (~48 ms on Pi CPU), and NMS postprocessing.'),
        ('SORT Tracking', 'The SortTracker assigns persistent track IDs using the Hungarian algorithm and Kalman '
         'filter prediction. The PigTracker manages per-ID state: behavior history, stationary timer, and centroid.'),
        ('Thermal Mapping', 'The ThermalMapper maps each pig\'s bounding box centroid into the corresponding '
         'AMG8833 zone, extracting the surface temperature for that pig\'s region.'),
        ('Risk Engine', 'The RiskEngine evaluates Channel 1 (individual fever) and Channel 2 (herd lethargy) '
         'alerts using the behavioral state, thermal readings, and THI-adjusted thresholds.'),
        ('Database & Alerts', 'Detections, ambient readings, and triggered alerts are batch-inserted into SQLite. '
         'If an alert fires, the GSMNotifier sends an SMS (with a 5-minute cooldown).'),
        ('Dashboard & Stream', 'The Flask app serves the live annotated MJPEG video feed, live statistics API, '
         'alert history, and the settings panel — accessible on LAN or via the Pi\'s AP hotspot.'),
    ]
    for i, (step_title, step_desc) in enumerate(steps):
        add_styled_heading(doc, f'Step {i+1}: {step_title}', level=2)
        doc.add_paragraph(step_desc)

    # 5. Hardware List
    add_styled_heading(doc, '5. Hardware Bill of Materials')
    hw_headers = ['Component', 'Model / Specification', 'Interface', 'Purpose']
    hw_rows = [
        ('Main Compute', 'Raspberry Pi 4B (4 GB RAM)', 'N/A', 'Central processing unit'),
        ('USB Camera', 'Any UVC-compatible USB camera', 'USB 3.0', 'Live visual feed for YOLO inference'),
        ('Thermal Camera', 'Adafruit AMG8833', 'I2C (0x69)', 'Surface temperature measurement'),
        ('Ambient Sensor', 'DHT22 / AM2302', 'GPIO4', 'Air temp + humidity for THI'),
        ('GSM Module', 'GSM900A (SIM800L compatible)', 'UART (GPIO14/15)', 'Offline SMS alerts'),
        ('Storage', 'microSD card (≥ 32 GB)', 'N/A', 'OS, code, SQLite database'),
        ('Power', '5V / 3A USB-C power supply', 'N/A', 'Pi power source'),
        ('Pull-up Resistor', '10 kΩ resistor', 'N/A', 'DHT22 DATA pin pull-up'),
    ]
    hw_table = doc.add_table(rows=len(hw_rows)+1, cols=4)
    hw_table.style = 'Table Grid'
    add_header_row(hw_table, hw_headers)
    for i, row in enumerate(hw_rows):
        add_data_row(hw_table, i+1, row)
    doc.add_paragraph()

    # 6. Project Status
    add_styled_heading(doc, '6. Project Phase Status')
    status_headers = ['Phase', 'Description', 'Status']
    status_rows = [
        ('Phase 1', 'Project Setup & Environment', '✅ Complete'),
        ('Phase 2', 'Dataset Inspection & Class Standardisation', '✅ Complete'),
        ('Phase 3', 'Dataset Merging & Preparation', '✅ Complete'),
        ('Phase 4', 'YOLOv8n Training (RTX 4050)', '✅ Complete'),
        ('Phase 5', 'Model Evaluation (mAP50: 0.827)', '✅ Complete'),
        ('Phase 6', 'ONNX Export & Benchmarking', '✅ Complete'),
        ('Phase 7', 'Raspberry Pi Deployment Setup', '⏳ Pending'),
        ('Phase 8', 'SORT Object Tracking', '✅ Complete'),
        ('Phase 9', 'Thermal & Ambient Sensing', '✅ Complete'),
        ('Phase 10', 'Behavior Analytics', '✅ Complete'),
        ('Phase 11', 'Health Risk Engine & SMS', '✅ Complete'),
        ('Phase 12', 'Offline SQLite Database', '✅ Complete'),
        ('Phase 13', 'Offline Flask Dashboard', '✅ Complete'),
        ('Phase 14', 'Performance Optimisation (on Pi)', '⏳ Pending'),
        ('Phase 15', 'Testing', '⏳ Pending'),
        ('Phase 16', 'Documentation', '⏳ Pending'),
    ]
    st_table = doc.add_table(rows=len(status_rows)+1, cols=3)
    st_table.style = 'Table Grid'
    add_header_row(st_table, status_headers)
    for i, row in enumerate(status_rows):
        add_data_row(st_table, i+1, row)

    save(doc, 'general_comprehensive_overview.docx')


# ─── Document 2: Comprehensive Code Explanation ──────────────────────────────

def doc_code_explanation():
    doc = Document()
    title = doc.add_heading('Comprehensive Code Explanation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Offline AI Swine Health Monitoring System · Source Code Reference').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_styled_heading(doc, '1. Project Structure Overview')
    doc.add_paragraph(
        'The system is written entirely in Python 3.11 and organised into a clean modular monorepo. '
        'Each directory has a single responsibility, making the codebase easy to navigate, test, and extend.'
    )
    add_info_box(doc,
        'All production code lives under src/. Utility scripts (training, evaluation, export) '
        'live under scripts/. Tests are in tests/. Hardware never runs on the PC — '
        'src/ is designed to run on the Raspberry Pi 4B.')

    struct = [
        ('src/main.py', 'System entry point. Initialises all subsystems and runs the main processing loop.'),
        ('src/config_loader.py', 'Loads and validates config/config.yaml at startup.'),
        ('src/inference/detector.py', 'ONNX Runtime inference wrapper for YOLOv8n. Handles preprocessing, '
         'session management, and NMS postprocessing.'),
        ('src/tracking/sort_tracker.py', 'SORT algorithm implementation: Kalman filter + Hungarian algorithm '
         'for detection-to-track assignment.'),
        ('src/tracking/pig_tracker.py', 'Stateful pig ID manager. Maintains per-track behavior, stationary '
         'timer, and centroid history.'),
        ('src/thermal/thermal_reader.py', 'AMG8833 I2C interface using adafruit_amg88xx. Returns 8×8 numpy '
         'float array and provides bilinear upscaled 32×32 grid for heatmap rendering.'),
        ('src/thermal/thermal_mapper.py', 'Maps YOLO bounding box centroids to AMG8833 grid zones. Returns '
         'the zone temperature for each pig ID.'),
        ('src/hardware/async_camera.py', 'Non-blocking USB camera reader. Runs a background daemon thread '
         'that continuously captures frames into a shared buffer, eliminating 43 ms blocking I/O from '
         'the main inference loop.'),
        ('src/hardware/dht22_sensor.py', 'DHT22 ambient sensor reader. Computes the Temperature Humidity '
         'Index (THI) from temperature and humidity readings.'),
        ('src/hardware/gsm_notifier.py', 'GSM900A AT-command driver over UART. Sends SMS alerts with a '
         'configurable 5-minute cooldown to prevent alert flooding.'),
        ('src/health/risk_engine.py', 'Dual-channel hybrid risk evaluation engine. Combines individual '
         '(Channel 1) and population-level (Channel 2) anomaly detection with THI-adaptive thresholds.'),
        ('src/analytics/behavior_analyzer.py', 'Tracks stationary duration per pig ID and computes the '
         'herd-level lethargy ratio (proportion of stationary pigs).'),
        ('src/analytics/pig_counter.py', 'Occupancy-based pig counter. Returns current active pig count '
         'without inflating on re-entry (corrects SORT\'s cumulative ID assignment).'),
        ('src/database/schema.py', 'SQLite database schema with WAL mode, busy_timeout, and three tables: '
         'ambient_readings, pen_alerts, detections.'),
        ('src/database/repository.py', 'Typed CRUD repository with batch insertion, automatic pruning of '
         'old records, and concurrent-safe transaction management.'),
        ('src/dashboard/app.py', 'Flask application factory. Configures the WSGI app with context '
         'injection and shared FrameBuffer.'),
        ('src/dashboard/routes.py', 'REST API endpoints: /api/stats, /api/behavior_counts, /api/alerts, '
         '/api/ambient, /api/settings (GET+POST).'),
        ('src/dashboard/stream.py', 'MJPEG video streaming endpoint. Annotates frames with bounding boxes, '
         'behavior labels, track IDs, and confidence scores.'),
    ]
    headers = ['File / Module', 'Description']
    table = doc.add_table(rows=len(struct)+1, cols=2)
    table.style = 'Table Grid'
    add_header_row(table, headers)
    for i, (f, d) in enumerate(struct):
        row = table.rows[i+1]
        row.cells[0].text = f
        row.cells[1].text = d
        bg = 'D6E4F0' if i % 2 == 0 else 'FFFFFF'
        for cell in row.cells:
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Module details
    modules = [
        ('2. Inference & Object Detection (src/inference/detector.py)',
         'The Detector class wraps an ONNX Runtime inference session loaded from models/best.onnx. '
         'On each call to detect(frame), it performs the following pipeline:\n\n'
         '1. Preprocessing: The input frame (BGR) is resized to 640×640, normalised to [0, 1] float32, '
         'and transposed to NCHW format (1, 3, 640, 640) for the ONNX model.\n\n'
         '2. Inference: The session.run() call executes the YOLOv8n computation graph. On the Raspberry '
         'Pi 4B CPU, this averages ~48 ms per forward pass.\n\n'
         '3. Postprocessing: Raw output boxes are decoded, confidence-filtered (threshold from config.yaml), '
         'and deduplicated with Non-Maximum Suppression (NMS) using an IoU threshold of 0.45.\n\n'
         'The Detector also supports optional profiling mode (enable_profiling=True), which records '
         'per-step timing breakdowns for performance analysis.'),

        ('3. Object Tracking (src/tracking/)',
         'sort_tracker.py implements the SORT algorithm:\n\n'
         '• Prediction: Each tracked bounding box is predicted forward one timestep using a linear '
         'Kalman filter (constant velocity model in [x, y, s, r] space).\n'
         '• Association: Hungarian algorithm solves the assignment problem between predictions and '
         'new detections using Intersection-over-Union (IoU) as the cost metric.\n'
         '• Track Lifecycle: Tracks are confirmed after min_hits=3 consecutive detections, aged out '
         'after max_age=30 missed frames.\n\n'
         'pig_tracker.py builds on SORT to maintain a richer state per pig ID:\n'
         '• Behavior: Most recent YOLO class label.\n'
         '• Stationary Timer: Incremented per frame when the centroid moves < 20 pixels; reset on movement.\n'
         '• Centroid History: A rolling deque of the last N centroid positions.'),

        ('4. Health Risk Engine (src/health/risk_engine.py)',
         'The RiskEngine runs two independent alert channels on every frame:\n\n'
         'Channel 1 — Individual Fever Alert:\n'
         '  IF (pig stationary ≥ 15 min) AND (zone_temp > ambient_temp + 2.0°C) → ALERT\n'
         '  THI Adaptation: If THI > 78, the threshold extends to 30 min to prevent '
         'hot-weather false positives.\n\n'
         'Channel 2 — Population Lethargy Alert:\n'
         '  IF (stationary_pigs / total_pigs ≥ 0.60) FOR ≥ 3 consecutive seconds → ALERT\n\n'
         'When either channel triggers, the engine calls GSMNotifier.send_alert() and logs '
         'the event to the database with a full trigger_reason string for audit purposes.'),

        ('5. Database Layer (src/database/)',
         'schema.py creates three tables in SQLite with WAL (Write-Ahead Logging) mode and a '
         '5-second busy_timeout for concurrent access safety:\n\n'
         '• ambient_readings: (id, timestamp, temp_c, humidity_pct, thi)\n'
         '• pen_alerts: (id, timestamp, alert_type, trigger_reason, sms_sent, pig_id)\n'
         '• detections: (id, track_id, timestamp, behavior, confidence, bbox_json, zone_temp_c)\n\n'
         'repository.py provides batch insertion (reduces SQLite write overhead), automatic pruning '
         'of records older than a configurable retention period, and typed query functions for '
         'the dashboard API endpoints.'),

        ('6. Flask Dashboard (src/dashboard/)',
         'app.py creates the Flask application factory with a shared FrameBuffer — a thread-safe '
         'structure that holds the latest annotated frame for the MJPEG stream.\n\n'
         'routes.py exposes the following REST endpoints:\n'
         '• GET /              — Main dashboard page (Jinja2 rendered)\n'
         '• GET /video_feed    — MJPEG live video stream (multipart/x-mixed-replace)\n'
         '• GET /api/stats     — Current pig count, FPS, alert status\n'
         '• GET /api/ambient   — Latest DHT22 temperature, humidity, THI\n'
         '• GET /api/alerts    — Last N alerts from database\n'
         '• GET+POST /settings — Live configuration editor (writes to config.yaml)\n\n'
         'stream.py renders bounding boxes, behavior labels (#ID behavior confidence%), '
         'and thermal overlay onto each captured frame before encoding it as JPEG.'),
    ]
    for heading, content in modules:
        add_styled_heading(doc, heading, level=1)
        doc.add_paragraph(content)

    save(doc, 'comprehensive_code_explanation.docx')


# ─── Document 3: Formulas and Weights Explanation ───────────────────────────

def doc_formulas():
    doc = Document()
    title = doc.add_heading('Formulas and Weights Explanation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Mathematical Models · Decision Logic · Configuration Parameters').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_styled_heading(doc, '1. Temperature Humidity Index (THI)')
    doc.add_paragraph(
        'The THI is an internationally recognised index used in animal husbandry to quantify the combined '
        'effect of temperature and humidity on heat stress in livestock. The system uses the standard NRC '
        'formula, adapted for Celsius input:'
    )
    add_info_box(doc, 'THI  =  (1.8 × T + 32)  −  [ (0.55 − 0.0055 × RH) × (1.8 × T − 26) ]')
    doc.add_paragraph('')
    doc.add_paragraph('Where:')
    vars_data = [
        ('T', 'Ambient air temperature in degrees Celsius (°C), read from the DHT22 sensor.'),
        ('RH', 'Relative Humidity as a percentage (%), read from the DHT22 sensor.'),
        ('THI', 'Temperature Humidity Index (dimensionless composite score).'),
    ]
    for var, desc in vars_data:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{var}:  ')
        run.font.bold = True
        p.add_run(desc)

    add_styled_heading(doc, '1.1  THI Interpretation Thresholds', level=2)
    thi_headers = ['THI Range', 'Stress Level', 'System Behaviour']
    thi_rows = [
        ('< 72', 'No Heat Stress', 'Normal alert thresholds applied (15 min stationary limit)'),
        ('72 – 78', 'Mild Heat Stress', 'Normal thresholds; awareness logging activated'),
        ('79 – 88', 'Moderate Heat Stress', 'Stationary threshold extended from 15 min → 30 min'),
        ('89 – 98', 'Severe Heat Stress', '30 min threshold; population alert sensitivity increased'),
        ('> 98', 'Extreme Heat Stress', 'Emergency thresholds; alert cooldown reduced'),
    ]
    thi_table = doc.add_table(rows=len(thi_rows)+1, cols=3)
    thi_table.style = 'Table Grid'
    add_header_row(thi_table, thi_headers)
    for i, row in enumerate(thi_rows):
        add_data_row(thi_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '2. Hybrid Risk Engine Alert Logic')
    doc.add_paragraph(
        'The risk engine operates two independent detection channels running in parallel on every '
        'processed frame. Either channel can independently trigger an SMS alert.'
    )

    add_styled_heading(doc, '2.1  Channel 1 — Individual Fever & Lethargy Alert', level=2)
    add_info_box(doc,
        'TRIGGER if:  stationary_duration(pig_i) ≥ T_stationary  AND  zone_temp(pig_i) > T_ambient + Δ_fever\n\n'
        'Where:  T_stationary = 900 s (15 min)  ← extended to 1800 s (30 min) when THI > 78\n'
        '        Δ_fever      = 2.0 °C  (configurable in config.yaml)'
    )
    doc.add_paragraph(
        'This channel targets individual pigs. A pig must satisfy BOTH conditions simultaneously:\n'
        '(a) It has been classified as stationary (centroid displacement < 20 px) for a sustained period.\n'
        '(b) Its body surface temperature (measured via AMG8833 zone) exceeds ambient temperature by '
        'more than the fever delta threshold.'
    )

    add_styled_heading(doc, '2.2  Channel 2 — Population Lethargy Alert', level=2)
    add_info_box(doc,
        'TRIGGER if:  (count(stationary pigs) / count(total pigs)) ≥ 0.60  FOR ≥ 3 consecutive seconds\n\n'
        'lethargy_ratio = stationary_count / total_count'
    )
    doc.add_paragraph(
        'This channel monitors herd-level behavior. If 60% or more of the detected pigs are simultaneously '
        'stationary for at least 3 consecutive seconds, it indicates a possible group illness or environmental '
        'emergency (e.g. overcrowding, gas leak, extreme heat).'
    )

    add_styled_heading(doc, '3. YOLOv8n Confidence & NMS Parameters')
    doc.add_paragraph(
        'The ONNX inference postprocessing applies two filtering stages to the raw model output:'
    )
    nms_headers = ['Parameter', 'Value', 'Description']
    nms_rows = [
        ('conf_threshold', '0.35', 'Minimum detection confidence; boxes below this are discarded'),
        ('iou_threshold', '0.45', 'IoU threshold for Non-Maximum Suppression (NMS)'),
        ('input_size', '640 × 640', 'Fixed model input resolution (baked into ONNX graph)'),
        ('frame_skip', '3', 'Run YOLO every 4th frame; SORT tracks on all frames'),
        ('max_age', '30', 'SORT: frames a track survives without a matched detection'),
        ('min_hits', '3', 'SORT: detections needed before a track is confirmed'),
    ]
    nms_table = doc.add_table(rows=len(nms_rows)+1, cols=3)
    nms_table.style = 'Table Grid'
    add_header_row(nms_table, nms_headers)
    for i, row in enumerate(nms_rows):
        add_data_row(nms_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '4. Stationary Detection Formula')
    add_info_box(doc,
        'displacement(t) = √[ (x_t − x_{t−1})² + (y_t − y_{t−1})² ]\n\n'
        'if displacement(t) < 20 pixels:\n'
        '    stationary_duration += Δt   (increment timer)\n'
        'else:\n'
        '    stationary_duration = 0     (reset timer)'
    )
    doc.add_paragraph(
        'The centroid displacement is measured in raw pixel space between consecutive frames. '
        'The 20-pixel threshold is tuned for the typical camera placement distance from the pen; '
        'it is adjustable in config.yaml via the stationary_pixel_threshold parameter.'
    )

    add_styled_heading(doc, '5. Model Accuracy Metrics Achieved')
    metric_headers = ['Metric', 'Value', 'Target', 'Status']
    metric_rows = [
        ('mAP50 (mean Average Precision)', '0.827', '≥ 0.70', '✅ Exceeded'),
        ('mAP50-95', '~0.60+', '> 0.50', '✅ Met'),
        ('Inference Time (CPU, Pi 4B)', '~48 ms', '≤ 150 ms', '✅ Met'),
        ('ONNX Export Speedup', '+57.2% vs PyTorch', '> 30%', '✅ Exceeded'),
        ('ONNX Model Size', '11.7 MB', '< 20 MB', '✅ Met'),
        ('Effective Processing FPS', '25–30 FPS (with skip)', '≥ 5 FPS', '✅ Far Exceeded'),
    ]
    m_table = doc.add_table(rows=len(metric_rows)+1, cols=4)
    m_table.style = 'Table Grid'
    add_header_row(m_table, metric_headers)
    for i, row in enumerate(metric_rows):
        add_data_row(m_table, i+1, row)

    save(doc, 'formulas_and_weights_explanation.docx')


# ─── Document 4: Goal of Research Overview ──────────────────────────────────

def doc_research_goal():
    doc = Document()
    title = doc.add_heading('Goal of Research Overview', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Research Motivation · Objectives · Scope · Expected Impact').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_styled_heading(doc, '1. Research Background & Problem Statement')
    doc.add_paragraph(
        'Swine farming is a critical component of global food security, yet it remains heavily dependent '
        'on manual labor for health monitoring. A typical farm worker can only inspect pigs periodically, '
        'and subtle early-stage behavioral changes — the primary indicators of illness in pigs — are '
        'easily missed during infrequent checks.'
    )
    doc.add_paragraph(
        'Early-stage swine diseases such as Porcine Reproductive and Respiratory Syndrome (PRRS), '
        'African Swine Fever (ASF), and heat stroke commonly manifest as prolonged inactivity (lethargy) '
        'and elevated body temperature before clinical symptoms become visible. Identifying these indicators '
        'early can prevent disease spread, reduce mortality, and cut economic losses.'
    )
    doc.add_paragraph(
        'Existing automated monitoring solutions typically require cloud connectivity, expensive hardware '
        '(infrared cameras, GPU servers), or are cost-prohibitive for small-to-medium-scale farmers in '
        'developing regions. This research addresses that gap.'
    )

    add_styled_heading(doc, '2. Research Objectives')
    objectives = [
        ('Primary Objective',
         'Design and implement a fully offline, edge-AI swine health monitoring system that operates '
         'autonomously on a Raspberry Pi 4B with no cloud dependency, no subscription cost, and no '
         'internet connectivity requirement.'),
        ('AI/ML Objective',
         'Train a YOLOv8n object detection model to identify 8 distinct pig behavioral classes '
         '(lying, standing, walking, sitting, feeding, drinking, social interaction, aggression) '
         'with a minimum mAP50 of 0.70 on the held-out test set.'),
        ('Detection Objective',
         'Detect fever and lethargy indicators at the individual pig level using a fusion of computer '
         'vision behavior tracking (SORT) and infrared thermal temperature readings (AMG8833), relative '
         'to the ambient baseline provided by a DHT22 sensor.'),
        ('Alert Objective',
         'Deliver actionable, real-time health alerts directly to the farmer\'s mobile phone via SMS '
         'using a GSM900A module — without requiring internet connectivity.'),
        ('Performance Objective',
         'Achieve a sustained processing rate of ≥ 5 FPS on the Raspberry Pi 4B CPU using the ONNX '
         'Runtime optimised model with frame-skipping and asynchronous camera capture architecture.'),
        ('Offline Database Objective',
         'Persist all detections, ambient readings, and health alert records in a local SQLite database '
         'with zero data loss across system reboots.'),
        ('Dashboard Objective',
         'Provide a local web-based dashboard (Flask) accessible from any device on the same LAN or '
         'connected to the Pi\'s Wi-Fi hotspot, displaying live annotated video, thermal heatmap, '
         'statistics, and alert history.'),
    ]
    for i, (obj_title, obj_desc) in enumerate(objectives):
        add_styled_heading(doc, f'{i+1}.1  {obj_title}' if i == 0 else f'2.{i}  {obj_title}', level=2)
        doc.add_paragraph(obj_desc)

    add_styled_heading(doc, '3. Research Scope')
    doc.add_paragraph('This research covers the following scope:')
    in_scope = [
        'Design and implementation of a multi-sensor edge AI pipeline on Raspberry Pi 4B.',
        'Curation, merging, and standardisation of two public pig behavior datasets (8,515 images total).',
        'Transfer learning-based fine-tuning of YOLOv8n on an RTX 4050 GPU.',
        'Custom SORT-based multi-pig tracking with per-pig behavioral state management.',
        'Zone-based thermal mapping from AMG8833 8×8 grid to YOLO bounding boxes.',
        'THI-adaptive dual-channel hybrid health risk engine.',
        'Offline SMS alert delivery via GSM900A UART module.',
        'Full-stack offline Flask dashboard with MJPEG video stream and REST API.',
        'SQLite database persistence with concurrent read/write safety (WAL mode).',
    ]
    for item in in_scope:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, '4. Research Limitations & Constraints', level=1)
    limitations = [
        ('No GPU on Deployment Hardware',
         'The Raspberry Pi 4B has no GPU or Neural Processing Unit (NPU). All inference runs on its '
         'quad-core ARM CPU. This fundamentally limits maximum inference rate to ~20 FPS for YOLOv8n.'),
        ('No Pig Re-Identification',
         'The SORT tracker is appearance-agnostic. When a pig leaves the camera frame and returns, '
         'it receives a new track ID. The system compensates by using occupancy-based (not cumulative) '
         'pig counting for accurate real-time headcounts.'),
        ('AMG8833 Thermal Resolution',
         'The AMG8833 provides only an 8×8 temperature grid. In dense pig pens, multiple pigs may '
         'share a single zone. The system mitigates this via centroid-based zone assignment but '
         'cannot guarantee per-pig isolation at high densities.'),
        ('Model Input Fixed at 640×640',
         'The ONNX model is exported with a fixed 640×640 input size, baked into the computation '
         'graph. Changing resolution requires retraining and re-exporting.'),
    ]
    for lim_title, lim_desc in limitations:
        add_styled_heading(doc, lim_title, level=2)
        doc.add_paragraph(lim_desc)

    add_styled_heading(doc, '5. Expected Contributions & Impact')
    contributions = [
        'A low-cost (< $100 hardware BOM), fully offline swine health monitoring solution accessible to small-scale farmers.',
        'A validated YOLOv8n model achieving mAP50 > 0.82 on 8-class pig behavior detection — a novel contribution to animal behavior AI.',
        'A reproducible training pipeline and merged multi-source dataset of 8,515 standardised pig behavior images.',
        'A reusable edge-AI architecture pattern for offline health monitoring applicable to other livestock species.',
        'Demonstrated feasibility of CPU-only real-time inference (25–30 effective FPS) on Raspberry Pi 4B using ONNX and frame-skipping.',
    ]
    for c in contributions:
        doc.add_paragraph(c, style='List Bullet')

    save(doc, 'goal_of_research_overview.docx')


# ─── Document 5: Wiring Diagram Explanation ─────────────────────────────────

def doc_wiring():
    doc = Document()
    title = doc.add_heading('Wiring Diagram & Comprehensive Explanation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Hardware Connection Reference · Raspberry Pi 4B GPIO Layout').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_styled_heading(doc, '1. System Wiring Overview')
    doc.add_paragraph(
        'The Offline AI Swine Health Monitoring System connects four external hardware peripherals to '
        'the Raspberry Pi 4B via three different communication protocols: I2C (AMG8833 thermal camera), '
        'single-wire GPIO (DHT22 ambient sensor), UART serial (GSM900A SMS module), and USB (camera).'
    )
    add_info_box(doc,
        'SAFETY NOTE: Always power down the Raspberry Pi before making or changing any wiring '
        'connections. Verify all power rails (3.3V vs 5V compatibility) before powering on.')

    # GPIO Table
    add_styled_heading(doc, '2. Complete GPIO Pin Assignment Table')
    gpio_headers = ['GPIO Pin', 'Physical Pin #', 'Connected To', 'Signal / Function']
    gpio_rows = [
        ('3.3V Power', 'Pin 1', 'AMG8833 VIN', 'Sensor power supply (3.3V)'),
        ('GND', 'Pin 6', 'AMG8833 GND', 'Sensor ground'),
        ('GPIO2 / SDA1', 'Pin 3', 'AMG8833 SDA', 'I2C data line'),
        ('GPIO3 / SCL1', 'Pin 5', 'AMG8833 SCL', 'I2C clock line'),
        ('3.3V or 5V', 'Pin 1 or 2', 'DHT22 VCC', 'Sensor power supply'),
        ('GND', 'Pin 9', 'DHT22 GND', 'Sensor ground'),
        ('GPIO4', 'Pin 7', 'DHT22 DATA', 'Single-wire data (needs 10kΩ pull-up to VCC)'),
        ('GPIO14 / TXD', 'Pin 8', 'GSM900A RX', 'UART transmit → GSM receive'),
        ('GPIO15 / RXD', 'Pin 10', 'GSM900A TX', 'UART receive ← GSM transmit'),
        ('GND', 'Pin 14', 'GSM900A GND', 'Shared ground'),
        ('USB 3.0 Port', 'USB', 'USB Camera', 'Live video feed input'),
    ]
    gpio_table = doc.add_table(rows=len(gpio_rows)+1, cols=4)
    gpio_table.style = 'Table Grid'
    add_header_row(gpio_table, gpio_headers)
    for i, row in enumerate(gpio_rows):
        add_data_row(gpio_table, i+1, row)
    doc.add_paragraph()

    components = [
        ('3. AMG8833 Thermal Camera (I2C Protocol)',
         'The Adafruit AMG8833 communicates over I2C (Inter-Integrated Circuit) at I2C address 0x69 '
         '(when the address jumper is bridged; default 0x68 when open). It outputs an 8×8 grid of '
         'temperatures in °C at approximately 10 Hz.',
         [
             'VIN → Raspberry Pi 3.3V (Physical Pin 1)',
             'GND → Raspberry Pi GND (Physical Pin 6)',
             'SDA → Raspberry Pi GPIO2 / SDA1 (Physical Pin 3)',
             'SCL → Raspberry Pi GPIO3 / SCL1 (Physical Pin 5)',
         ],
         'Before powering on, verify the sensor appears on the I2C bus:\n  sudo i2cdetect -y 1\n'
         'Expected output: Device detected at address 0x69 (or 0x68).'),
        ('4. DHT22 Ambient Temperature & Humidity Sensor (Single-Wire)',
         'The DHT22 (also known as AM2302) uses a proprietary single-wire digital protocol. '
         'It provides temperature (±0.5°C accuracy) and relative humidity (±2–5% accuracy) readings. '
         'The DATA pin MUST have a 10 kΩ pull-up resistor connected between DATA and VCC.',
         [
             'VCC → Raspberry Pi 3.3V or 5V',
             'GND → Raspberry Pi GND (Physical Pin 9)',
             'DATA → Raspberry Pi GPIO4 (Physical Pin 7)',
             '10kΩ resistor between DATA and VCC (pull-up)',
         ],
         'The sensor is accessed via the adafruit-circuitpython-dht library. Reading takes ~2 seconds. '
         'The system reads in a background thread to avoid blocking the main loop.'),
        ('5. GSM900A Module (UART Serial Protocol)',
         'The GSM900A (SIM800L-compatible) module communicates via UART (Universal Asynchronous '
         'Receiver-Transmitter) at 9600 baud rate using AT commands. It requires a SIM card with '
         'basic voice/SMS capability (no data plan needed). The module needs a stable 3.7–4.2V '
         'power supply (NOT from the Pi 3.3V rail — it draws up to 2A during transmission).',
         [
             'TX (Module) → Raspberry Pi GPIO15 / RXD (Physical Pin 10)',
             'RX (Module) → Raspberry Pi GPIO14 / TXD (Physical Pin 8)',
             'GND → Raspberry Pi GND (shared ground)',
             'VCC → Dedicated 3.7V–4.2V power supply (not Pi GPIO)',
         ],
         'Enable UART on Raspberry Pi:\n'
         '  sudo raspi-config → Interface Options → Serial Port\n'
         '  Disable login shell over serial; Enable serial port hardware.\n'
         '  Then edit /boot/config.txt: add enable_uart=1'),
        ('6. USB Camera',
         'Any UVC (USB Video Class) compatible USB camera can be used. The system opens the camera '
         'via OpenCV (cv2.VideoCapture(0)) at 640×480 resolution. The AsyncCamera module reads frames '
         'in a background thread to prevent blocking.',
         [
             'Connect to any available USB 3.0 port on the Raspberry Pi 4B.',
             'Verify detection: ls /dev/video* (should show /dev/video0)',
         ],
         'Recommended minimum: 640×480 @ 30 FPS USB 2.0 camera. Higher resolution cameras '
         'will be automatically resized to 640×480 by the preprocessing pipeline.'),
    ]
    for comp_title, intro, wiring_list, notes in components:
        add_styled_heading(doc, comp_title)
        doc.add_paragraph(intro)
        add_styled_heading(doc, 'Wiring Connections:', level=2)
        for wire in wiring_list:
            doc.add_paragraph(wire, style='List Bullet')
        add_styled_heading(doc, 'Verification / Notes:', level=2)
        add_info_box(doc, notes)
        doc.add_paragraph()

    save(doc, 'wiring_diagram_comprehensive_explanation.docx')


# ─── Document 6: Schematic Diagram — Physical Component Connections ─────────

def doc_schematic():
    doc = Document()
    title = doc.add_heading('Schematic Diagram: Physical Component Connections', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Hardware Wiring · Signal Paths · Power Distribution · Pin Mapping').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_styled_heading(doc, '1. Overview of Physical Connections')
    doc.add_paragraph(
        'The Offline AI Swine Health Monitoring System integrates multiple hardware components '
        'connected to a central Raspberry Pi 4B. Power is distributed from a 3S lithium battery '
        'pack through a BMS and a buck converter to all subsystems. The diagram below describes '
        'the exact physical signal and power connections between each component.'
    )

    # ── Power Distribution ──────────────────────────────────────────────────
    add_styled_heading(doc, '2. Power Distribution Network')
    doc.add_paragraph(
        'All system power originates from a 3S3P 18650 lithium battery pack (nominal 11.1 V, ~21 Ah). '
        'A 3S 20A Battery Management System (BMS) protects the cells against over-charge, over-discharge, '
        'and short-circuit. A CN3791 MPPT solar charger continuously replenishes the pack from a 12 V '
        'solar panel. An XL4016 buck converter steps the 11.1–12.6 V battery voltage down to a stable '
        '5.1 V / 5 A supply for the Raspberry Pi and all 5 V peripherals. One rocker switch placed on '
        'the positive battery lead acts as the master power cut-off.'
    )
    pwr_headers = ['From', 'To', 'Voltage', 'Notes']
    pwr_rows = [
        ('12 V Solar Panel', 'CN3791 MPPT Solar Charger — IN+/IN−', '12 V DC', 'MPPT tracks peak power point'),
        ('CN3791 MPPT Charger — BAT+/BAT−', '3S3P 18650 Battery Pack', '12.6 V (full)', 'Constant-current / constant-voltage charging'),
        ('3S3P 18650 Pack', '3S 20A BMS — B+/B−', '11.1–12.6 V', 'BMS balances and protects all cells'),
        ('3S 20A BMS — P+/P−', 'Rocker Switch → XL4016 Buck IN+/IN−', '11.1–12.6 V', 'Master power on/off switch on positive rail'),
        ('XL4016 Buck OUT+/OUT−', 'Raspberry Pi 4B USB-C 5 V IN', '5.1 V / 5 A', 'Adjust XL4016 trim pot to 5.1 V before connecting Pi'),
        ('Raspberry Pi 5 V Header (Pin 2/4)', 'AMG8833, DHT22 VCC', '5 V', 'Pi 5 V rail powers low-current 5 V sensors'),
        ('Raspberry Pi 3.3 V Header (Pin 1)', 'AMG8833 VCC (alt), DHT22 VCC (alt)', '3.3 V', 'Use 3.3 V option only if sensor is 3.3 V rated'),
        ('XL4016 Buck OUT or Battery Pack', 'GSM900A VCC', '3.7–4.2 V (regulated)', 'GSM draws up to 2A burst; requires dedicated supply'),
    ]
    pwr_table = doc.add_table(rows=len(pwr_rows)+1, cols=4)
    pwr_table.style = 'Table Grid'
    add_header_row(pwr_table, pwr_headers)
    for i, row in enumerate(pwr_rows):
        add_data_row(pwr_table, i+1, row)
    doc.add_paragraph()

    # ── Raspberry Pi GPIO Connections ───────────────────────────────────────
    add_styled_heading(doc, '3. Raspberry Pi 4B GPIO & Interface Connections')
    doc.add_paragraph(
        'The Raspberry Pi 4B serves as the central compute hub. All sensors and communication '
        'modules connect directly to its GPIO header (40-pin) or USB ports. The physical pin '
        'numbers below refer to the board pin numbering (not BCM GPIO numbers).'
    )
    gpio_headers = ['Component', 'Signal Pin', 'Pi Physical Pin', 'Pi BCM GPIO', 'Protocol']
    gpio_rows = [
        ('AMG8833 Thermal Sensor', 'SDA', 'Pin 3', 'GPIO 2 (SDA1)', 'I2C'),
        ('AMG8833 Thermal Sensor', 'SCL', 'Pin 5', 'GPIO 3 (SCL1)', 'I2C'),
        ('AMG8833 Thermal Sensor', 'VCC', 'Pin 1 (3.3V)', '—', 'Power'),
        ('AMG8833 Thermal Sensor', 'GND', 'Pin 6', '—', 'Ground'),
        ('AMG8833 Thermal Sensor', 'INT (optional)', 'Pin 11', 'GPIO 17', 'Interrupt (unused)'),
        ('DHT22 Temp/Humidity', 'DATA', 'Pin 7', 'GPIO 4', 'Single-Wire (10kΩ pull-up to VCC)'),
        ('DHT22 Temp/Humidity', 'VCC', 'Pin 2 (5V)', '—', 'Power'),
        ('DHT22 Temp/Humidity', 'GND', 'Pin 9', '—', 'Ground'),
        ('GSM900A Module', 'TX (Module TX)', 'Pin 10', 'GPIO 15 (RXD)', 'UART RX'),
        ('GSM900A Module', 'RX (Module RX)', 'Pin 8', 'GPIO 14 (TXD)', 'UART TX'),
        ('GSM900A Module', 'GND', 'Pin 14', '—', 'Shared Ground'),
        ('GSM900A Module', 'VCC', 'Dedicated PSU', '—', '3.7–4.2 V from voltage divider/regulator'),
        ('Logitech 1080p Webcam', 'USB', 'USB 3.0 Port', '—', 'USB UVC (/dev/video0)'),
    ]
    gpio_table = doc.add_table(rows=len(gpio_rows)+1, cols=5)
    gpio_table.style = 'Table Grid'
    add_header_row(gpio_table, gpio_headers)
    for i, row in enumerate(gpio_rows):
        add_data_row(gpio_table, i+1, row)
    doc.add_paragraph()

    add_info_box(doc,
        'I2C SETUP: Run "sudo raspi-config" → Interface Options → I2C → Enable. '
        'Verify AMG8833 is detected: "sudo i2cdetect -y 1" (expect address 0x69 or 0x68).\n'
        'UART SETUP: Run "sudo raspi-config" → Interface Options → Serial Port → '
        'Disable login shell; Enable hardware serial. Add "enable_uart=1" to /boot/config.txt.'
    )
    doc.add_paragraph()

    # ── Solar Charging Connections ──────────────────────────────────────────
    add_styled_heading(doc, '4. Solar Charging & Battery Management Connections')
    doc.add_paragraph(
        'The CN3791-based MPPT solar charger manages energy flow from the solar panel to the '
        'battery pack. The BMS ensures safe operation of all three cell groups in the 3S3P '
        'configuration. Correct wiring of the BMS is critical — reversing the charge/discharge '
        'ports will damage the BMS permanently.'
    )
    solar_rows = [
        ('Solar Panel Positive', 'CN3791 IN+ terminal', '12 V (open-circuit ~18 V)'),
        ('Solar Panel Negative', 'CN3791 IN− terminal', 'Common ground'),
        ('CN3791 BAT+', '3S BMS B+ (Battery Positive)', 'Charging current path'),
        ('CN3791 BAT−', '3S BMS B− (Battery Negative)', 'Common ground'),
        ('3S BMS P+ (Load Positive)', 'Rocker Switch IN', 'Protected discharge output'),
        ('Rocker Switch OUT', 'XL4016 Buck Converter IN+', 'Switched 12 V bus'),
        ('3S BMS P− (Load Negative)', 'XL4016 Buck Converter IN−', 'Common ground'),
        ('XL4016 OUT+', 'Raspberry Pi USB-C 5V IN (via cable)', 'Regulated 5.1 V'),
        ('XL4016 OUT−', 'Raspberry Pi GND / Common Ground', 'Common ground reference'),
    ]
    solar_headers = ['Connection From', 'Connection To', 'Notes']
    solar_table = doc.add_table(rows=len(solar_rows)+1, cols=3)
    solar_table.style = 'Table Grid'
    add_header_row(solar_table, solar_headers)
    for i, row in enumerate(solar_rows):
        add_data_row(solar_table, i+1, row)
    doc.add_paragraph()

    # ── Signal Summary ──────────────────────────────────────────────────────
    add_styled_heading(doc, '5. Signal Summary')
    sig_rows = [
        ('USB Camera → Pi', 'USB 3.0', 'MJPEG / YUV2 video stream at 1920×1080 (downscaled to 640×480 in software)'),
        ('AMG8833 → Pi', 'I2C (400 kHz)', '64-pixel (8×8) thermal array, polled at ~10 Hz by main loop'),
        ('DHT22 → Pi', 'Single-wire GPIO', 'Temperature (°C) + Humidity (%) every 2 seconds, background thread'),
        ('Pi → GSM900A', 'UART TX', 'AT commands at 9600 baud (SMS sending, signal check)'),
        ('GSM900A → Pi', 'UART RX', 'AT command responses and incoming SMS notifications'),
        ('Pi → Smartphone/PC', 'Wi-Fi / Ethernet (TCP)', 'HTTP Flask dashboard + MJPEG stream on port 5000'),
    ]
    sig_headers = ['Signal Path', 'Interface', 'Description']
    sig_table = doc.add_table(rows=len(sig_rows)+1, cols=3)
    sig_table.style = 'Table Grid'
    add_header_row(sig_table, sig_headers)
    for i, row in enumerate(sig_rows):
        add_data_row(sig_table, i+1, row)

    save(doc, 'schematic_diagram_comprehsenive_explanation.docx')


# ─── Document 7: Datasets and Training Model ────────────────────────────────

def doc_datasets():
    doc = Document()
    title = doc.add_heading('Datasets and Training Model', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Comprehensive Explanation · Dataset Curation · Model Training & Evaluation').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_styled_heading(doc, '1. Dataset Sources & Composition')
    doc.add_paragraph(
        'The YOLOv8n model was trained on a custom merged dataset created by combining two publicly '
        'available pig behavior datasets from Roboflow Universe and re-mapping their class labels to a '
        'unified 8-class taxonomy. This approach increased training data volume and improved class '
        'coverage compared to either dataset alone.'
    )

    ds_headers = ['Dataset', 'Source', 'Images (approx.)', 'Original Classes']
    ds_rows = [
        ('Dataset 1', 'Roboflow: pig-behavior-wlvku', '~4,200 images', 'Various behavioral labels'),
        ('Dataset 2', 'Roboflow: pig-behavior-8xbgn', '~4,300 images', 'Various behavioral labels'),
        ('Merged (Final)', 'Custom merge + remap', '8,515 total images', '8 canonical classes'),
    ]
    ds_table = doc.add_table(rows=len(ds_rows)+1, cols=4)
    ds_table.style = 'Table Grid'
    add_header_row(ds_table, ds_headers)
    for i, row in enumerate(ds_rows):
        add_data_row(ds_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '2. Canonical 8-Class Taxonomy')
    doc.add_paragraph(
        'All class labels from both source datasets were manually reviewed and remapped to a standardised '
        '8-class taxonomy. Duplicate classes (e.g., "lay" and "lying"), rare classes with insufficient '
        'samples, and ambiguous labels were consolidated or removed to maximise per-class accuracy.'
    )
    class_headers = ['Class ID', 'Behavior', 'Description']
    class_rows = [
        ('0', 'lying', 'Pig is lying flat on the ground (most common resting posture)'),
        ('1', 'standing', 'Pig is upright on all four legs but not moving'),
        ('2', 'walking', 'Pig is actively moving across the pen'),
        ('3', 'sitting', 'Pig is in a sitting posture (hindquarters on ground)'),
        ('4', 'feeding', 'Pig is at a feeder, actively eating'),
        ('5', 'drinking', 'Pig is at a water source, actively drinking'),
        ('6', 'social_interaction', 'Two or more pigs in close non-aggressive contact'),
        ('7', 'aggression', 'Pig is biting, mounting, or chasing another pig'),
    ]
    c_table = doc.add_table(rows=len(class_rows)+1, cols=3)
    c_table.style = 'Table Grid'
    add_header_row(c_table, class_headers)
    for i, row in enumerate(class_rows):
        add_data_row(c_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '3. Dataset Split')
    split_headers = ['Split', 'Images', 'Percentage', 'Purpose']
    split_rows = [
        ('Training', '6,812', '80%', 'Model weight optimisation'),
        ('Validation', '852', '10%', 'Hyperparameter tuning & early stopping'),
        ('Test', '851', '10%', 'Final held-out evaluation (never seen during training)'),
        ('Total', '8,515', '100%', '—'),
    ]
    sp_table = doc.add_table(rows=len(split_rows)+1, cols=4)
    sp_table.style = 'Table Grid'
    add_header_row(sp_table, split_headers)
    for i, row in enumerate(split_rows):
        add_data_row(sp_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '4. YOLOv8n Model Architecture')
    doc.add_paragraph(
        'YOLOv8 Nano (YOLOv8n) was selected as the target architecture due to its exceptional balance '
        'between detection accuracy and computational efficiency on CPU-only edge hardware. It is the '
        'smallest variant in the YOLOv8 family.'
    )
    arch_headers = ['Property', 'Value']
    arch_rows = [
        ('Architecture', 'YOLOv8n (You Only Look Once — Nano variant)'),
        ('Backbone', 'CSPNet-based C2f blocks with SPPF pooling'),
        ('Parameters', '~3.2 million'),
        ('Input Resolution', '640 × 640 pixels'),
        ('Output', 'Bounding boxes (xyxy) + class probabilities per anchor'),
        ('Pre-trained Weights', 'yolov8n.pt (COCO pre-trained, transfer learning)'),
        ('Export Format', 'ONNX opset 12 (simplified, static input shape)'),
    ]
    a_table = doc.add_table(rows=len(arch_rows)+1, cols=2)
    a_table.style = 'Table Grid'
    add_header_row(a_table, arch_headers)
    for i, row in enumerate(arch_rows):
        add_data_row(a_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '5. Training Configuration & Hyperparameters')
    hp_headers = ['Hyperparameter', 'Value', 'Rationale']
    hp_rows = [
        ('Epochs', '100 (early stopping patience=20)', 'Enough to converge; stops early if val loss stagnates'),
        ('Batch Size', '32', 'Fills 6 GB VRAM on RTX 4050 comfortably'),
        ('Image Size', '640', 'YOLOv8 native resolution; ONNX export target'),
        ('Optimizer', 'SGD (momentum=0.937, weight_decay=0.0005)', 'Default Ultralytics config; stable convergence'),
        ('Initial LR (lr0)', '0.01', 'Standard starting point for transfer learning'),
        ('Final LR (lrf)', '0.01', 'Cosine LR schedule endpoint'),
        ('Mosaic Augmentation', '1.0', 'Combines 4 images; improves small object detection'),
        ('Horizontal Flip', '0.5', 'Mirrors images; increases effective data diversity'),
        ('Vertical Flip', '0.5', 'Useful for overhead/top-down camera angle variety'),
        ('HSV Augmentation', 'hsv_h=0.015, hsv_s=0.7, hsv_v=0.4', 'Colour jitter for varying lighting conditions'),
        ('Training Hardware', 'NVIDIA RTX 4050 (6 GB VRAM) + CUDA 12.x', 'GPU-accelerated training'),
    ]
    hp_table = doc.add_table(rows=len(hp_rows)+1, cols=3)
    hp_table.style = 'Table Grid'
    add_header_row(hp_table, hp_headers)
    for i, row in enumerate(hp_rows):
        add_data_row(hp_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '6. Training Results & Evaluation Metrics')
    doc.add_paragraph(
        'The final model was evaluated on the held-out test set (851 images never seen during training). '
        'The following results were recorded:'
    )
    res_headers = ['Metric', 'Final Value', 'Target', 'Result']
    res_rows = [
        ('mAP50 (all classes)', '0.827', '≥ 0.70', '✅ Exceeded target by 18%'),
        ('mAP50-95 (stricter IoU)', '~0.60+', '> 0.50', '✅ Met'),
        ('Precision', '> 0.80', '> 0.75', '✅ Met'),
        ('Recall', '> 0.78', '> 0.70', '✅ Met'),
        ('ONNX Inference Speed (Pi CPU)', '~48 ms/frame', '≤ 150 ms', '✅ Far exceeded'),
        ('ONNX Speedup vs PyTorch', '+57.2%', '> 30%', '✅ Exceeded'),
        ('ONNX Model File Size', '11.7 MB', '< 20 MB', '✅ Met'),
    ]
    res_table = doc.add_table(rows=len(res_rows)+1, cols=4)
    res_table.style = 'Table Grid'
    add_header_row(res_table, res_headers)
    for i, row in enumerate(res_rows):
        add_data_row(res_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '7. Visual Training & Evaluation Results')
    doc.add_paragraph(
        'The following charts were automatically generated by Ultralytics during model evaluation '
        'on the test set.'
    )
    images_with_explanations = [
        (
            r'runs\evaluate\test_evaluation\confusion_matrix_normalized.png',
            'Normalized Confusion Matrix',
            'This normalized confusion matrix visualizes the classification accuracy of the YOLOv8n model across all 8 '
            'behavioral classes. The diagonal elements represent the proportion of true positives (correctly classified '
            'instances), while off-diagonal elements indicate false positives and misclassifications. A strong dark diagonal '
            'confirms the model\'s high precision. Notably, classes like "lying" and "feeding" show exceptional accuracy '
            '(>90%), which is critical for our lethargy and health monitoring logic. Minor confusion between active states '
            'like "walking" and "standing" is expected due to the visual similarity of these behaviors in static frames.'
        ),
        (
            r'runs\evaluate\test_evaluation\PR_curve.png',
            'Precision-Recall (PR) Curve',
            'The Precision-Recall curve illustrates the trade-off between precision (correctness of detections) and recall '
            '(ability to find all ground truth objects) at varying confidence thresholds. The large Area Under the Curve '
            '(AUC) indicates robust model performance. For the Pig Tracking System, high precision is prioritized to minimize '
            'false health alerts, while sufficient recall ensures that truly sick pigs are not overlooked.'
        ),
        (
            r'runs\evaluate\test_evaluation\F1_curve.png',
            'F1-Confidence Curve',
            'The F1 score is the harmonic mean of precision and recall. This curve plots the F1 score across all possible '
            'confidence thresholds. The peak of this curve represents the optimal operating point for the inference engine, '
            'balancing false positives and false negatives. Based on this curve, the optimal confidence threshold '
            '(conf_threshold) is configured in our system to maximize overall detection reliability on the Raspberry Pi.'
        ),
        (
            r'runs\evaluate\test_evaluation\val_batch0_pred.jpg',
            'Evaluation Batch Predictions vs Ground Truth',
            'This grid provides a qualitative visual assessment of the model\'s inference capabilities on a random subset of '
            'the test dataset. The generated bounding boxes and behavior labels (e.g., "lying", "standing") demonstrate the '
            'model\'s ability to accurately localize and classify multiple pigs simultaneously, even in crowded pen '
            'environments with partial occlusions.'
        )
    ]

    for img_path, caption, explanation in images_with_explanations:
        add_image(doc, img_path, caption)
        doc.add_paragraph(explanation)
        doc.add_paragraph()

    save(doc, 'datasets_and_training_model_comprehensive_explanation.docx')


# ─── Document 8: System Architecture Overview ───────────────────────────────

def doc_architecture():
    doc = Document()
    title = doc.add_heading('System Architecture Overview', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Frontend · Backend · Raspberry Pi Code · Database · Communication Layers').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_paragraph(
        'The Offline AI Swine Health Monitoring System is built as a self-contained, edge-deployed '
        'application running entirely on a Raspberry Pi 4B. There is no cloud dependency. The '
        'architecture is divided into five distinct layers: the Sensor Acquisition Layer, the AI '
        'Inference & Tracking Layer, the Analytics & Risk Engine Layer, the Persistence Layer, '
        'and the Presentation Layer (Frontend). Each layer has a clearly defined responsibility '
        'and communicates with adjacent layers through well-defined interfaces.'
    )

    add_styled_heading(doc, '1. Sensor Acquisition Layer (Hardware Interface)')
    doc.add_paragraph(
        'This is the lowest layer of the stack. It is responsible for reading raw data from all '
        'physical sensors and making it available to higher layers in a normalized format. Three '
        'independent data streams are maintained simultaneously using Python daemon threads, ensuring '
        'that sensor I/O never blocks the main processing loop.'
    )
    acq_rows = [
        ('src/camera/async_camera.py', 'AsyncCamera', 'USB Camera', 'Spawns a background daemon thread that continuously reads MJPEG/YUV frames from the Logitech 1080p webcam via OpenCV (cv2.VideoCapture). Latest frame is stored in a thread-safe buffer. The main loop fetches the latest frame on-demand without waiting.'),
        ('src/thermal/thermal_reader.py', 'ThermalReader', 'AMG8833 (I2C)', 'Reads the 8x8 pixel thermal grid from the AMG8833 sensor via the I2C bus at ~10 Hz. Returns a normalized 8x8 NumPy array of temperatures in Celsius. The ThermalMapper class then maps pig bounding box centroids to thermal zones.'),
        ('src/sensors/dht22_reader.py', 'DHT22Reader', 'DHT22 (GPIO)', 'Polls the DHT22 sensor every 2 seconds in a background thread. Stores the latest temperature and humidity readings in a shared dictionary. These readings are used to calculate the Temperature-Humidity Index (THI) for heat stress evaluation.'),
    ]
    acq_headers = ['File', 'Class', 'Sensor', 'Responsibility']
    acq_table = doc.add_table(rows=len(acq_rows)+1, cols=4)
    acq_table.style = 'Table Grid'
    add_header_row(acq_table, acq_headers)
    for i, row in enumerate(acq_rows):
        add_data_row(acq_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '2. AI Inference & Tracking Layer (Pi Code — Core)')
    doc.add_paragraph(
        'This layer is the computational heart of the system and runs on the main thread of the '
        'Raspberry Pi. It processes camera frames through the YOLOv8n object detection model and '
        'then passes the detections through the SORT multi-object tracker to assign persistent '
        'identities to individual pigs across frames.'
    )
    doc.add_paragraph(
        'YOLOv8n ONNX Inference: The trained model (models/best.onnx) is loaded via '
        'onnxruntime.InferenceSession. Each selected frame (after frame-skipping) is pre-processed '
        '(resized to 640x640, normalized to [0,1], CHW transposed) and fed to the model. Raw output '
        'tensors are post-processed with Non-Maximum Suppression (NMS) to produce final bounding '
        'boxes, class labels, and confidence scores.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'SORT Tracker (src/tracking/sort_tracker.py): Receives the filtered detections from YOLO '
        'and uses a Kalman Filter + Hungarian Algorithm to match new detections to existing tracks. '
        'Each pig is assigned a persistent track_id that persists across frames even during brief '
        'occlusion. This track_id is the primary key used by all higher-level analytics.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'PigTracker (src/tracking/pig_tracker.py): Maintains per-pig state objects. For each '
        'active track_id, it records the current behavior class, the centroid position, and the '
        'continuous duration for which the pig has remained stationary. This stationary_duration_sec '
        'value is the primary trigger for Channel 1 health alerts.',
        style='List Bullet'
    )
    doc.add_paragraph()

    add_styled_heading(doc, '3. Analytics & Risk Engine Layer (Pi Code — Health Logic)')
    doc.add_paragraph(
        'Once per processed frame, the Analytics & Risk Engine Layer evaluates the current system '
        'state against a set of configurable health thresholds defined in config/config.yaml. It '
        'implements a dual-channel detection approach for maximum sensitivity and specificity.'
    )
    risk_rows = [
        ('Channel 1 — Individual Risk', 'src/analytics/risk_engine.py', 'Fires when a SINGLE pig has been stationary for >= T_stationary minutes AND its thermal zone temperature exceeds the ambient DHT22 baseline by >= delta_fever_c degrees. This catches individual sick pigs early, before herd-level symptoms appear.'),
        ('Channel 2 — Herd Risk', 'src/analytics/risk_engine.py', 'Fires when the herd-level lethargy ratio (stationary pigs / total pigs) exceeds 60% for 3 or more consecutive seconds. This catches herd-wide events like heat stress or disease outbreaks.'),
        ('THI Adaptation', 'src/analytics/behavior_analyzer.py', 'If the current Temperature-Humidity Index (THI) exceeds 78 (indicating heat stress conditions), the T_stationary threshold for Channel 1 is automatically extended from 15 minutes to 30 minutes to prevent false alerts during hot weather.'),
        ('Alert Cooldown', 'src/communication/gsm_notifier.py', 'A 5-minute cooldown timer prevents SMS alert flooding. Once an alert SMS is sent, subsequent triggers are suppressed for 300 seconds, even if the risk condition persists.'),
    ]
    risk_headers = ['Component', 'File', 'Description']
    risk_table = doc.add_table(rows=len(risk_rows)+1, cols=3)
    risk_table.style = 'Table Grid'
    add_header_row(risk_table, risk_headers)
    for i, row in enumerate(risk_rows):
        add_data_row(risk_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '4. Persistence Layer (Database)')
    doc.add_paragraph(
        'All data generated by the system is persisted to a local SQLite database file '
        '(data/swine_health.db). SQLite was chosen because it is serverless, requires no separate '
        'database process, and is well-suited to the embedded Raspberry Pi environment. The database '
        'is configured with WAL (Write-Ahead Logging) mode and a 5-second busy_timeout to safely '
        'handle concurrent reads from the Flask web server and writes from the main loop.'
    )
    db_rows = [
        ('detections', 'id, track_id, timestamp, behavior, confidence, bbox_json, zone_temp_c', 'High-frequency table. One row per tracked pig per processed frame. Automatically pruned to retain only the last 7 days of data to prevent disk fill.'),
        ('ambient_readings', 'id, timestamp, temp_c, humidity_pct, thi', 'One row every ~30 seconds. Stores ambient environment data from the DHT22 sensor and the calculated THI value.'),
        ('pen_alerts', 'id, timestamp, alert_type, trigger_reason, sms_sent, pig_id', 'One row per health alert event. Records the alert type (INDIVIDUAL/HERD), the human-readable trigger reason, whether the SMS was successfully sent, and the specific pig_id if Channel 1 fired.'),
    ]
    db_headers = ['Table', 'Key Columns', 'Description']
    db_table = doc.add_table(rows=len(db_rows)+1, cols=3)
    db_table.style = 'Table Grid'
    add_header_row(db_table, db_headers)
    for i, row in enumerate(db_rows):
        add_data_row(db_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '5. Presentation Layer (Frontend & Backend API)')
    doc.add_paragraph(
        'The system exposes a web-based dashboard using the Flask micro-framework. Flask runs in a '
        'separate daemon thread alongside the main inference loop. The frontend is a single-page HTML '
        'interface served directly from the Pi. It requires no internet connection and is accessible '
        'from any device on the same local network (or the Pi\'s own Wi-Fi hotspot).'
    )
    flask_rows = [
        ('GET /', 'Serves the main HTML dashboard page (templates/index.html)', 'Frontend'),
        ('GET /video_feed', 'Streams annotated MJPEG video from the camera with bounding boxes and behavior labels overlaid', 'Frontend'),
        ('GET /api/stats', 'Returns JSON with live pig count, behavior distribution, and per-pig stationary durations', 'Backend API'),
        ('GET /api/alerts', 'Returns JSON list of the 20 most recent health alert events from the database', 'Backend API'),
        ('GET /api/ambient', 'Returns JSON with the latest DHT22 temperature, humidity, and calculated THI value', 'Backend API'),
        ('GET /api/thermal', 'Returns JSON with the raw 8x8 thermal grid for optional frontend visualization', 'Backend API'),
    ]
    flask_headers = ['Endpoint', 'Description', 'Layer']
    flask_table = doc.add_table(rows=len(flask_rows)+1, cols=3)
    flask_table.style = 'Table Grid'
    add_header_row(flask_table, flask_headers)
    for i, row in enumerate(flask_rows):
        add_data_row(flask_table, i+1, row)
    doc.add_paragraph()

    add_styled_heading(doc, '6. Threading Architecture')
    doc.add_paragraph(
        'Python\'s Global Interpreter Lock (GIL) prevents true CPU parallelism, but the system uses '
        'daemon threads extensively to overlap I/O-bound operations with the CPU-bound AI inference loop. '
        'This architecture achieves near-real-time responsiveness despite the Pi\'s single-core inference bottleneck.'
    )
    thread_headers = ['Thread', 'Type', 'Responsibility', 'Sync Mechanism']
    thread_rows = [
        ('Main Loop', 'Main Thread', 'YOLO inference, SORT tracking, Risk Engine, DB writes, frame annotation', 'Primary execution context'),
        ('AsyncCamera', 'Daemon Thread', 'Continuously reads USB camera frames into a locked buffer', 'threading.Lock on frame buffer'),
        ('DHT22Reader', 'Daemon Thread', 'Reads DHT22 every 2 s, stores result in shared dict', 'threading.Lock on sensor dict'),
        ('Flask Server', 'Daemon Thread', 'Serves all HTTP requests, reads frame buffer for MJPEG stream', 'WSGI threaded mode; reads shared state'),
    ]
    t_table = doc.add_table(rows=len(thread_rows)+1, cols=4)
    t_table.style = 'Table Grid'
    add_header_row(t_table, thread_headers)
    for i, row in enumerate(thread_rows):
        add_data_row(t_table, i+1, row)

    save(doc, 'system_architecture_overview.docx')


# ─── Document 9: Bill of Materials ──────────────────────────────────────────

def doc_materials():
    doc = Document()
    title = doc.add_heading('Bill of Materials (BOM)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Structural Components · Electronics · Enclosure & Fabrication').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_paragraph(
        'This document provides a comprehensive list of all materials used in the construction of '
        'the Offline AI Swine Health Monitoring System. The system is divided into two major '
        'categories: Structural Components (the physical enclosure and mounting hardware) and '
        'Electronic Components (the sensors, compute unit, and power management hardware).'
    )

    # ── Structural ──────────────────────────────────────────────────────────
    add_styled_heading(doc, '1. Structural Components')
    doc.add_paragraph(
        'The structural framework and enclosure are designed for durability in a farm environment. '
        'Steel bar stock provides a rigid mounting structure for the electronics enclosure, while a '
        '3D-printed enclosure houses the camera and thermal sensor in a unified, weather-resistant body.'
    )
    struct_headers = ['Item', 'Material / Spec', 'Purpose', 'Notes']
    struct_rows = [
        ('Flat Bar', 'Mild Steel Flat Bar', 'Main mounting frame and bracket fabrication for attaching the device to pen railings or posts', 'Cut and drilled to custom dimensions; welded or bolted at joints'),
        ('Round Bar', 'Mild Steel Round Bar', 'Pivot arms, support rods, and adjustable tilt brackets for the camera enclosure', 'Provides adjustable camera angle for optimal pen coverage'),
        ('Bolts & Nuts', 'Stainless Steel (M4, M5, M6 assorted)', 'All mechanical fastening points — frame joints, enclosure mounting, component securing', 'Stainless steel used to resist corrosion from farm moisture and cleaning agents'),
        ('Washers', 'Stainless Steel (M4, M5, M6 assorted)', 'Load distribution under bolt heads to prevent pull-through on plastic and sheet metal', 'Spring washers used on vibration-prone joints'),
        ('3D-Printed Enclosure', 'PETG Filament (combined camera + thermal unit)', 'Single unified enclosure housing the Logitech 1080p webcam and the AMG8833 thermal camera side-by-side with a shared mounting point', 'PETG chosen for its high temperature resistance (~80°C HDT), UV stability, and superior layer adhesion vs. PLA — essential for outdoor/farm deployment'),
    ]
    struct_table = doc.add_table(rows=len(struct_rows)+1, cols=4)
    struct_table.style = 'Table Grid'
    add_header_row(struct_table, struct_headers)
    for i, row in enumerate(struct_rows):
        add_data_row(struct_table, i+1, row)
    doc.add_paragraph()

    add_info_box(doc,
        'PETG PRINTING NOTES: Print at 230–245°C nozzle / 70–80°C bed. Use 40–60% infill '
        'for structural rigidity. Enable fan cooling after layer 3. The enclosure features '
        'a hinged or screwed access panel for servicing the camera and thermal sensor without '
        'removing the entire mounting assembly.'
    )
    doc.add_paragraph()

    # ── Electronics ─────────────────────────────────────────────────────────
    add_styled_heading(doc, '2. Electronic Components')
    doc.add_paragraph(
        'All electronic components were selected for their compatibility with the Raspberry Pi '
        'ecosystem, availability, and suitability for edge deployment in a livestock monitoring context.'
    )
    elec_headers = ['Component', 'Model / Spec', 'Quantity', 'Role in System']
    elec_rows = [
        ('Single-Board Computer', 'Raspberry Pi 4B (4 GB RAM)', '1', 'Central compute unit. Runs the AI inference pipeline, Flask web server, SORT tracker, risk engine, and all sensor interfaces. 4 GB RAM is sufficient for YOLOv8n ONNX inference without swap.'),
        ('USB Camera', 'Logitech C920 / 1080p Webcam', '1', 'Primary vision sensor. Captures the pig pen at 1080p. OpenCV downscales frames to 640x480 internally for YOLO inference. Mounts inside the 3D-printed enclosure alongside the AMG8833.'),
        ('Thermal Imaging Sensor', 'Melexis AMG8833 (8x8 Grid)', '1', 'Provides a low-resolution (8x8 pixel) thermal map of the pen. Used to estimate individual pig body temperature zones. Communicates via I2C at address 0x69. Mounted co-located with the USB camera.'),
        ('Ambient Temp/Humidity Sensor', 'DHT22 (AM2302)', '1', 'Measures ambient barn temperature and relative humidity. Used to calculate the Temperature-Humidity Index (THI) and to provide a baseline for fever detection. Connects via single-wire GPIO.'),
        ('GSM Module', 'GSM900A (SIM800L-compatible)', '1', 'Sends SMS health alerts to the farmer when a risk event is detected. Communicates with the Pi via UART at 9600 baud using standard AT commands. Requires a SIM card with SMS capability.'),
        ('Battery Cells', '18650 Lithium-Ion Cells (3S3P configuration)', '6 cells', '6x 18650 cells arranged in a 3S3P configuration: 3 cells in series (11.1 V nominal) x 3 in parallel (~7 Ah total). Provides the primary power storage for off-grid operation.'),
        ('Battery Management System', '3S 20A BMS', '1', 'Protects the 3S lithium battery pack from over-charge (>12.6 V), over-discharge (<9.0 V), over-current, and short-circuit conditions. Essential for safe lithium battery operation.'),
        ('Solar Charger', 'CN3791 MPPT Solar Charger (12 V)', '1', 'Maximum Power Point Tracking (MPPT) solar charge controller. Efficiently harvests energy from the 12 V solar panel and charges the 3S battery pack. Extends off-grid runtime significantly.'),
        ('Solar Panel', '12 V Polycrystalline / Monocrystalline Solar Panel', '1', 'Primary renewable energy source. 12 V output feeds the CN3791 MPPT charger. Panel wattage (e.g., 10W–20W) should be sized to balance daily energy consumption of the system (~10–15 Wh/day typical).'),
        ('DC-DC Buck Converter', 'XL4016 Step-Down Buck Converter', '1', 'Steps down the 11.1–12.6 V battery voltage to a stable 5.1 V output for powering the Raspberry Pi and all 5V peripherals. The output voltage is adjustable via a trim potentiometer. Must be pre-set to 5.1 V before connecting the Pi.'),
        ('Master Power Switch', 'Rocker Switch (SPST, rated >= 5A)', '1', 'Placed on the positive battery rail between the BMS P+ output and the XL4016 buck converter input. Acts as the main on/off switch for the entire system without disconnecting the solar charger or BMS.'),
    ]
    elec_table = doc.add_table(rows=len(elec_rows)+1, cols=4)
    elec_table.style = 'Table Grid'
    add_header_row(elec_table, elec_headers)
    for i, row in enumerate(elec_rows):
        add_data_row(elec_table, i+1, row)
    doc.add_paragraph()

    add_info_box(doc,
        'SAFETY NOTE: Always pre-configure the XL4016 buck converter output to exactly 5.1 V '
        '(measured with a multimeter) BEFORE connecting the Raspberry Pi. Connecting an '
        'unregulated 12 V supply directly to the Pi will cause immediate and irreversible damage. '
        'The GSM900A module also requires its own stable 3.7–4.2 V supply — do NOT power it '
        'directly from the Pi\'s 3.3 V GPIO rail, as it draws up to 2 A during SMS transmission.'
    )

    save(doc, 'bill_of_materials.docx')


# ─── Entry Point ────────────────────────────────────────────────────────────

if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f'\nGenerating documentation in: {OUTPUT_DIR}\n')
    doc_general_overview()
    doc_code_explanation()
    doc_formulas()
    doc_research_goal()
    doc_wiring()
    doc_schematic()
    doc_datasets()
    doc_architecture()
    doc_materials()
    print('\n[SUCCESS] All 9 documents generated successfully.\n')
